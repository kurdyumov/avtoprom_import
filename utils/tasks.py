import json
import os.path
import re
from urllib.parse import quote_plus
import pythoncom
import requests
import win32com.client
from docx import Document
from docx2pdf import convert
from sqlalchemy import create_engine, text
import datetime
from utils.do_predict import do_forecast
from utils.tg import tg_bot


class Tasks:
    def __init__(self, app_path: str):
        self.app_path = app_path

    def prepare_doc(self, source_config, header_repl, table_repl, par_repl):
        report = Document(os.path.join(self.app_path, 'reports', 'templates', 'state_query_result.docx'))
        for section in report.sections:
            header = section.header
            for p in header.paragraphs:
                for key, value in header_repl.items():
                    p.text = p.text.replace(key, value)
        for p in report.paragraphs:
            for key, value in par_repl.items():
                # print(p.text, str(key), str(value))
                p.text = p.text.replace(str(key), str(value))
        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in table_repl.items():
                            p.text = p.text.replace(str(key), str(value))
        return report

    def write_pdf(self, doc, title):
        message = None
        report_name = title.replace(" ", "_")
        save_path = os.path.join(self.app_path, 'reports', report_name)
        doc.save(save_path)
        pythoncom.CoInitialize()
        try:
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False

            pdf_path = save_path.replace('.docx', '.pdf')
            w2p = word.Documents.Open(os.path.abspath(save_path))
            w2p.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            w2p.Close(False)
            word.Quit()
            # convert(save_path, save_path.replace('.docx', '.pdf'))
            # os.remove(save_path)
            return message, pdf_path
        except Exception as e:
            print(f"Ошибка при конвертации в PDF: {str(e)}")
        finally:
            os.remove(save_path)
            pythoncom.CoUninitialize()

        try:
            os.remove(save_path)
        except:
            pass

    def pred_and_print(self, source_key, source_config, params, error_msg):
        title = source_config['title']

        models_config = os.path.join(self.app_path, "data", "json", "models.json")
        with open(models_config, 'r', encoding='utf-8') as f:
            data = json.load(f)
        model = data['models'][source_config['model']]

        if error_msg:
            pred = error_msg
            params = 'Ошибка'
        else:
            try:
                pred = do_forecast(
                    source_config['model'],
                    params,
                    self.app_path
                )
                obj = {model['fields'][k]: v for k, v in params.items()}
                params = json.dumps(obj, sort_keys=True, indent=4, ensure_ascii=False)
            except Exception as e:
                pred = f'Ошибка в ходе выполнения прогноза: {str(e)}'

        header = {
            '{DATETIME}': datetime.datetime.now().strftime('%d %B %Y г. %H:%M:%S')
        }

        model_name = source_config['model']
        m_type = {
            'classification': 'классификация',
            'regression': 'регрессия',
            'ranking': 'ранжирование'
        }
        content = {
            '{TASK_NUM}': source_key,
            '{PROCESS}': source_config['title'],
            '{MODEL}': f"{model_name}.cbm ({m_type[model['type']]})",
            '{TARGET}': model['target'],
            '{RESULT}': pred,
            '{PARAMS}': params
        }
        paragraph = {
            '{TASK_NAME}': title,
            '{TASK_NUM}': source_key,
        }

        report = self.prepare_doc(source_config, header_repl=header, table_repl=content, par_repl=paragraph)
        base_name = f'{title}_{source_key}_{datetime.datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx'
        message, url = self.write_pdf(report, base_name)
        try:
            message = f"Объект: {source_config['title']}\nЗадача: {model['target']}\nПрогноз: {pred}"
            tg_bot().broadcast_report(url, source_key, message)
        except Exception as e:
            print(f'Не удалось разослать: {str(e)}')

    def parse_sql(self, source_key, source_config):
        try:
            source = source_config['source']['db']
            driver = source['driver']
            if source['driver'] == 'mysql':
                driver = f'{source["driver"]}+pymysql'
            connection_path = f'{driver}://{source["user"]}:{quote_plus(source["password"])}@{source["host"]}:{source["port"]}/{source["dbname"]}'
            engine = create_engine(connection_path, echo=False)
            with engine.connect() as c:
                # source["query"] - SQL-запрос, его надо проверить
                res = c.execute(text(source["query"])).fetchone()
                params = {k: res[v] for k, v in source_config['fields'].items()}
                print(params)
                self.pred_and_print(source_key, source_config, params, None)
        except Exception as e:
            err = f'Не удалось проанализировать данные из СУБД: {str(e)}'
            print(err)
            self.pred_and_print(source_key, source_config, None, err)

    def parse_web(self, source_key, source_config):
        try:
            req = requests.get(source_config['source']['host'])
            params = req.json()
            kv = {v: k for k, v in source_config['fields'].items()}
            params = {kv[k]: v for k, v in params.items()}
            self.pred_and_print(source_key, source_config, params, None)
        except Exception as e:
            err = f'Не удалось проанализировать данные из API: {str(e)}'
            print(err)
            self.pred_and_print(source_key, source_config, None, err)

    # def parse_web(self, source_key, source_config):
    #     title = source_config['title']
    #     models_config = os.path.join(self.app_path, "data", "json", "models.json")
    #     with open(models_config, 'r', encoding='utf-8') as f:
    #         data = json.load(f)
    #     model = data['models'][source_config['model']]
    #
    #     try:
    #         print('должен отредачить')
    #         req = requests.get(source_config['source']['host'])
    #         params = req.json()
    #         kv = {v: k for k, v in source_config['fields'].items()}
    #         params = {kv[k]: v for k, v in params.items()}
    #         pred = do_forecast(
    #             source_config['model'],
    #             params,
    #             self.app_path
    #         )
    #         obj = {model['fields'][k]: v for k, v in params.items()}
    #         params = json.dumps(obj, sort_keys=True, indent=4, ensure_ascii=False)
    #     except Exception as e:
    #         pred = f'Не удалось проанализировать данные из API: {e}'
    #         params = 'Ошибка'
    #
    #     header = {
    #         '{DATETIME}': datetime.datetime.now().strftime('%d %B %Y г. %H:%M:%S')
    #     }
    #
    #     content = {
    #         '{TASK_NUM}': source_key,
    #         '{PROCESS}': source_config['title'],
    #         '{MODEL}': f"{source_config['model']}.cbm",
    #         '{TARGET}': model['target'],
    #         '{RESULT}': pred,
    #         '{PARAMS}': params
    #     }
    #     paragraph = {
    #         '{TASK_NAME}': title,
    #         '{TASK_NUM}': source_key,
    #     }
    #
    #     report = self.prepare_doc(source_config, header_repl=header, table_repl=content, par_repl=paragraph)
    #     base_name = f'{title}_{source_key}_{datetime.datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx'
    #     message, url = self.write_pdf(report, base_name)
    #     try:
    #         message = f"Объект: {source_config['title']}\nЗадача: {model['target']}\nПрогноз: {pred}"
    #         tg_bot().broadcast_report(url, source_key, message)
    #     except Exception as e:
    #         print(f'[parse_web] conn: {e}; url: {url}')

    # def parse_sql(self, source_key, source_config):
    #     title = source_config['title']
    #     header = {
    #         '{DATETIME}': datetime.datetime.now().strftime('%d %B %Y г. %H:%M:%S')
    #     }
    #
    #     models_config = os.path.join(self.app_path, "data", "json", "models.json")
    #     with open(models_config, 'r', encoding='utf-8') as f:
    #         data = json.load(f)
    #     model = data['models'][source_config['model']]
    #
    #     try:
    #         source = source_config['source']['db']
    #         driver = source['driver']
    #         if source['driver'] == 'mysql':
    #             driver = f'{source["driver"]}+pymysql'
    #         connection_path = f'{driver}://{source["user"]}:{quote_plus(source["password"])}@{source["host"]}:{source["port"]}/{source["dbname"]}'
    #         engine = create_engine(connection_path, echo=False)
    #         with engine.connect() as c:
    #             res = c.execute(text(source["query"])).fetchone()
    #             params = {k: res[v] for k, v in source_config['fields'].items()}
    #             print(f'sql params: {params}')
    #
    #             # Предсказание
    #             pred = do_forecast(
    #                 source_config['model'],
    #                 params,
    #                 self.app_path
    #             )
    #             params = json.dumps(params, sort_keys=True, indent=4)
    #     except Exception as e:
    #         pred = f'Не удалось проанализировать данные из БД: {e}'
    #         params = 'Ошибка'
    #
    #     content = {
    #         '{TASK_NUM}': source_key,
    #         '{PROCESS}': source_config['title'],
    #         '{MODEL}': f"{source_config['model']}.cbm",
    #         '{TARGET}': model['target'],
    #         '{RESULT}': pred,
    #         '{PARAMS}': params
    #     }
    #     paragraph = {
    #         '{TASK_NAME}': title,
    #         '{TASK_NUM}': source_key,
    #     }
    #
    #     report = self.prepare_doc(source_config, header_repl=header, table_repl=content, par_repl=paragraph)
    #     base_name = f'{title}_{source_key}_{datetime.datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx'
    #     message, url = self.write_pdf(report, base_name)
    #     try:
    #         message = f"Объект: {source_config['title']}\nЗадача: {model['target']}\nПрогноз: {pred}"
    #         tg_bot().broadcast_report(url, message)
    #     except:
    #         pass
    #     print(f'[parse_sql] conn: {source_key}')
